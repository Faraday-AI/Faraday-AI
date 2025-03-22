from typing import Dict, List, Any, Optional, BinaryIO
import logging
import tempfile
import os
from pathlib import Path
import magic
import json
from PIL import Image
import docx
import PyPDF2
import csv
import pandas as pd
import numpy as np
from io import StringIO, BytesIO

logger = logging.getLogger(__name__)

class FileProcessingService:
    def __init__(self):
        self.supported_formats = {
            'text': ['.txt', '.md', '.json', '.csv', '.xml'],
            'document': ['.doc', '.docx', '.pdf', '.rtf'],
            'image': ['.jpg', '.jpeg', '.png', '.gif', '.bmp'],
            'spreadsheet': ['.xls', '.xlsx', '.csv'],
            'code': ['.py', '.js', '.html', '.css', '.java', '.cpp']
        }

    async def analyze_file(
        self,
        file: BinaryIO,
        filename: str
    ) -> Dict[str, Any]:
        """Analyze a file and extract relevant information."""
        try:
            # Get file type using python-magic
            file_content = file.read()
            mime_type = magic.from_buffer(file_content, mime=True)
            file.seek(0)  # Reset file pointer

            # Get file extension
            ext = Path(filename).suffix.lower()

            # Basic file info
            file_info = {
                "filename": filename,
                "mime_type": mime_type,
                "extension": ext,
                "size": len(file_content),
                "type": self._get_file_type(ext)
            }

            # Process based on file type
            if ext in self.supported_formats['text'] + self.supported_formats['code']:
                content_info = await self._process_text_file(file_content)
            elif ext in self.supported_formats['document']:
                content_info = await self._process_document(file_content, ext)
            elif ext in self.supported_formats['image']:
                content_info = await self._process_image(file_content)
            elif ext in self.supported_formats['spreadsheet']:
                content_info = await self._process_spreadsheet(file_content)
            else:
                content_info = {"warning": "Unsupported file format for detailed analysis"}

            return {
                "status": "success",
                "file_info": file_info,
                "content_info": content_info
            }
        except Exception as e:
            logger.error(f"Error analyzing file: {str(e)}")
            return {"status": "error", "error": str(e)}

    def _get_file_type(self, extension: str) -> str:
        """Determine file type from extension."""
        for file_type, extensions in self.supported_formats.items():
            if extension in extensions:
                return file_type
        return "unknown"

    async def _process_text_file(
        self,
        content: bytes
    ) -> Dict[str, Any]:
        """Process text-based files."""
        try:
            text = content.decode('utf-8')
            lines = text.split('\n')
            words = text.split()

            # Basic text analysis
            analysis = {
                "line_count": len(lines),
                "word_count": len(words),
                "char_count": len(text),
                "avg_line_length": np.mean([len(line) for line in lines]) if lines else 0,
                "avg_word_length": np.mean([len(word) for word in words]) if words else 0
            }

            # Try to detect if it's JSON
            try:
                json.loads(text)
                analysis["format"] = "json"
            except json.JSONDecodeError:
                analysis["format"] = "plain_text"

            return analysis
        except Exception as e:
            logger.error(f"Error processing text file: {str(e)}")
            return {"error": str(e)}

    async def _process_document(
        self,
        content: bytes,
        extension: str
    ) -> Dict[str, Any]:
        """Process document files (PDF, DOCX)."""
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=extension) as temp_file:
                temp_file.write(content)
                temp_file.flush()

                if extension == '.pdf':
                    with open(temp_file.name, 'rb') as pdf_file:
                        pdf_reader = PyPDF2.PdfReader(pdf_file)
                        analysis = {
                            "page_count": len(pdf_reader.pages),
                            "metadata": pdf_reader.metadata if pdf_reader.metadata else {},
                            "text_preview": pdf_reader.pages[0].extract_text()[:1000]
                        }
                elif extension in ['.doc', '.docx']:
                    doc = docx.Document(temp_file.name)
                    analysis = {
                        "paragraph_count": len(doc.paragraphs),
                        "section_count": len(doc.sections),
                        "text_preview": '\n'.join(p.text for p in doc.paragraphs[:5])
                    }

            os.unlink(temp_file.name)
            return analysis
        except Exception as e:
            logger.error(f"Error processing document: {str(e)}")
            return {"error": str(e)}

    async def _process_image(
        self,
        content: bytes
    ) -> Dict[str, Any]:
        """Process image files."""
        try:
            image = Image.open(BytesIO(content))
            analysis = {
                "format": image.format,
                "mode": image.mode,
                "size": image.size,
                "width": image.width,
                "height": image.height,
                "aspect_ratio": image.width / image.height,
                "dpi": image.info.get('dpi', 'Unknown')
            }

            # Basic color analysis for RGB images
            if image.mode == 'RGB':
                img_array = np.array(image)
                analysis.update({
                    "mean_rgb": img_array.mean(axis=(0, 1)).tolist(),
                    "std_rgb": img_array.std(axis=(0, 1)).tolist()
                })

            return analysis
        except Exception as e:
            logger.error(f"Error processing image: {str(e)}")
            return {"error": str(e)}

    async def _process_spreadsheet(
        self,
        content: bytes
    ) -> Dict[str, Any]:
        """Process spreadsheet files."""
        try:
            # Try reading as CSV first
            try:
                df = pd.read_csv(BytesIO(content))
            except:
                # If not CSV, try Excel
                df = pd.read_excel(BytesIO(content))

            analysis = {
                "row_count": len(df),
                "column_count": len(df.columns),
                "columns": df.columns.tolist(),
                "data_types": df.dtypes.astype(str).to_dict(),
                "null_counts": df.isnull().sum().to_dict(),
                "numeric_columns": df.select_dtypes(include=[np.number]).columns.tolist(),
                "preview": df.head().to_dict() if len(df) > 0 else {}
            }

            # Basic statistical analysis for numeric columns
            if len(analysis["numeric_columns"]) > 0:
                stats = df[analysis["numeric_columns"]].describe()
                analysis["statistics"] = stats.to_dict()

            return analysis
        except Exception as e:
            logger.error(f"Error processing spreadsheet: {str(e)}")
            return {"error": str(e)}

    async def extract_text(
        self,
        file: BinaryIO,
        filename: str
    ) -> Dict[str, Any]:
        """Extract text content from various file formats."""
        try:
            ext = Path(filename).suffix.lower()
            content = file.read()

            if ext in self.supported_formats['text'] + self.supported_formats['code']:
                text = content.decode('utf-8')
            elif ext == '.pdf':
                with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as temp_file:
                    temp_file.write(content)
                    temp_file.flush()
                    with open(temp_file.name, 'rb') as pdf_file:
                        pdf_reader = PyPDF2.PdfReader(pdf_file)
                        text = '\n'.join(page.extract_text() for page in pdf_reader.pages)
                    os.unlink(temp_file.name)
            elif ext in ['.doc', '.docx']:
                with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as temp_file:
                    temp_file.write(content)
                    temp_file.flush()
                    doc = docx.Document(temp_file.name)
                    text = '\n'.join(paragraph.text for paragraph in doc.paragraphs)
                    os.unlink(temp_file.name)
            else:
                return {
                    "status": "error",
                    "error": "Unsupported file format for text extraction"
                }

            return {
                "status": "success",
                "text": text,
                "char_count": len(text),
                "word_count": len(text.split())
            }
        except Exception as e:
            logger.error(f"Error extracting text: {str(e)}")
            return {"status": "error", "error": str(e)} 
