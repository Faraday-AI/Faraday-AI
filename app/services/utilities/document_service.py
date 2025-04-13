from typing import Dict, Any
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import logging
from datetime import datetime
from app.models.lesson_plan import LessonPlan
from app.core.config import get_settings
from functools import lru_cache

logger = logging.getLogger(__name__)
settings = get_settings()

class DocumentService:
    def __init__(self):
        self.template_dir = settings.TEMPLATE_DIR
        self.temp_dir = settings.TEMP_DIR
        
        # Create directories if they don't exist
        os.makedirs(self.template_dir, exist_ok=True)
        os.makedirs(self.temp_dir, exist_ok=True)

    def _create_styles(self, doc: Document):
        """Create custom styles for the document."""
        # Header style
        style = doc.styles.add_style('Custom Heading', WD_STYLE_TYPE.PARAGRAPH)
        style.font.size = Pt(16)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 51, 102)
        style.paragraph_format.space_after = Pt(12)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subheading style
        style = doc.styles.add_style('Custom Subheading', WD_STYLE_TYPE.PARAGRAPH)
        style.font.size = Pt(14)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 51, 102)
        style.paragraph_format.space_after = Pt(6)

        # Section style
        style = doc.styles.add_style('Section Header', WD_STYLE_TYPE.PARAGRAPH)
        style.font.size = Pt(12)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    def _add_smart_goal_section(self, doc: Document, objective):
        """Add SMART goal details in a structured format."""
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        row = table.rows[0].cells
        row[0].text = "SMART Goal Components"
        row[1].text = "Description"
        
        for component in ['specific', 'measurable', 'achievable', 'relevant', 'time_bound']:
            row = table.add_row().cells
            row[0].text = component.title()
            row[1].text = getattr(objective.smart_goal, component)

    async def generate_lesson_plan(self, lesson_plan: LessonPlan) -> str:
        """Generate a lesson plan document with enhanced formatting."""
        try:
            doc = Document()
            self._create_styles(doc)
            
            # Header
            header = doc.add_paragraph('Elizabeth Public Schools', style='Custom Heading')
            subheader = doc.add_paragraph('Physical Education and Health Department', style='Custom Heading')
            
            # Basic Information Table
            info_table = doc.add_table(rows=1, cols=2)
            info_table.style = 'Table Grid'
            
            # Add basic information
            cells = info_table.rows[0].cells
            cells[0].text = f"Teacher: {lesson_plan.teacher_name}"
            cells[1].text = f"Week of: {lesson_plan.week_of}"
            
            row = info_table.add_row().cells
            row[0].text = f"Subject: {lesson_plan.subject}"
            row[1].text = f"Grade Level: {lesson_plan.grade_level}"
            
            row = info_table.add_row().cells
            row[0].text = f"Unit: {lesson_plan.unit_title}"
            row[1].text = f"Lesson: {lesson_plan.lesson_title}"
            
            row = info_table.add_row().cells
            row[0].text = f"Date: {lesson_plan.date}"
            row[1].text = f"Duration: {lesson_plan.duration} minutes"
            
            # Standards
            doc.add_paragraph('Standards', style='Custom Subheading')
            for standard in lesson_plan.standards:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"{standard.code} ({standard.type}): ").bold = True
                p.add_run(standard.description)
            
            # Objectives and SMART Goals
            doc.add_paragraph('Learning Objectives', style='Custom Subheading')
            for obj in lesson_plan.objectives:
                p = doc.add_paragraph(style='List Number')
                p.add_run("Objective: ").bold = True
                p.add_run(obj.description)
                
                if obj.language_objective:
                    p = doc.add_paragraph()
                    p.add_run("Language Objective: ").bold = True
                    p.add_run(obj.language_objective)
                
                self._add_smart_goal_section(doc, obj)
            
            # Essential Question
            doc.add_paragraph('Essential Question', style='Section Header')
            doc.add_paragraph(lesson_plan.essential_question)
            
            # Do Now
            doc.add_paragraph('Do Now', style='Section Header')
            doc.add_paragraph(lesson_plan.do_now)
            
            # Materials
            doc.add_paragraph('Materials Needed', style='Section Header')
            for material in lesson_plan.materials_needed:
                doc.add_paragraph(material, style='List Bullet')
            
            # Instructional Plan
            doc.add_paragraph('Instructional Plan', style='Custom Subheading')
            
            # Anticipatory Set
            p = doc.add_paragraph(style='Section Header')
            p.add_run("Anticipatory Set")
            doc.add_paragraph(lesson_plan.anticipatory_set)
            
            # Direct Instruction
            p = doc.add_paragraph(style='Section Header')
            p.add_run("Direct Instruction")
            doc.add_paragraph(lesson_plan.direct_instruction)
            
            # Guided Practice
            p = doc.add_paragraph(style='Section Header')
            p.add_run("Guided Practice")
            for activity in lesson_plan.guided_practice:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"{activity.name} ({activity.duration} minutes)").bold = True
                doc.add_paragraph(activity.description)
                if activity.materials:
                    doc.add_paragraph(f"Materials: {', '.join(activity.materials)}")
                if activity.modifications:
                    doc.add_paragraph("Modifications:")
                    for group, mod in activity.modifications.items():
                        doc.add_paragraph(f"- {group.upper()}: {mod}", style='List Bullet')
            
            # Independent Practice
            p = doc.add_paragraph(style='Section Header')
            p.add_run("Independent Practice")
            for activity in lesson_plan.independent_practice:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"{activity.name} ({activity.duration} minutes)").bold = True
                doc.add_paragraph(activity.description)
            
            # Closure
            p = doc.add_paragraph(style='Section Header')
            p.add_run("Closure")
            doc.add_paragraph(lesson_plan.closure)
            
            # Assessment
            doc.add_paragraph('Assessment', style='Custom Subheading')
            for assessment in lesson_plan.assessments:
                p = doc.add_paragraph()
                p.add_run(f"{assessment.type}: ").bold = True
                p.add_run(assessment.description)
                if assessment.criteria:
                    doc.add_paragraph("Success Criteria:")
                    for criterion in assessment.criteria:
                        doc.add_paragraph(criterion, style='List Bullet')
                if assessment.modifications:
                    doc.add_paragraph("Assessment Modifications:")
                    for group, mod in assessment.modifications.items():
                        doc.add_paragraph(f"- {group.upper()}: {mod}", style='List Bullet')
            
            # Differentiation
            doc.add_paragraph('Differentiation Plan', style='Custom Subheading')
            
            # ELL Strategies
            p = doc.add_paragraph(style='Section Header')
            p.add_run("ELL Strategies")
            for key, value in lesson_plan.differentiation.ell_strategies.items():
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"{key.replace('_', ' ').title()}: ").bold = True
                p.add_run(value)
            
            # IEP Accommodations
            p = doc.add_paragraph(style='Section Header')
            p.add_run("IEP Accommodations")
            for key, value in lesson_plan.differentiation.iep_accommodations.items():
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"{key.title()}: ").bold = True
                p.add_run(value)
            
            # 504 Accommodations
            p = doc.add_paragraph(style='Section Header')
            p.add_run("504 Accommodations")
            for key, value in lesson_plan.differentiation.section_504_accommodations.items():
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"{key.title()}: ").bold = True
                p.add_run(value)
            
            # Gifted and Talented
            p = doc.add_paragraph(style='Section Header')
            p.add_run("Gifted and Talented Enrichment")
            for key, value in lesson_plan.differentiation.gifted_talented_enrichment.items():
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"{key.title()}: ").bold = True
                p.add_run(value)
            
            # Additional Components
            if any([lesson_plan.homework, lesson_plan.notes, lesson_plan.reflection, lesson_plan.next_steps]):
                doc.add_paragraph('Additional Notes', style='Custom Subheading')
                
                if lesson_plan.homework:
                    p = doc.add_paragraph()
                    p.add_run("Homework: ").bold = True
                    p.add_run(lesson_plan.homework)
                
                if lesson_plan.notes:
                    p = doc.add_paragraph()
                    p.add_run("Notes: ").bold = True
                    p.add_run(lesson_plan.notes)
                
                if lesson_plan.reflection:
                    p = doc.add_paragraph()
                    p.add_run("Reflection: ").bold = True
                    p.add_run(lesson_plan.reflection)
                
                if lesson_plan.next_steps:
                    p = doc.add_paragraph()
                    p.add_run("Next Steps: ").bold = True
                    p.add_run(lesson_plan.next_steps)
            
            # Save document
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"lesson_plan_{timestamp}.docx"
            filepath = os.path.join(self.temp_dir, filename)
            doc.save(filepath)
            
            return filepath
            
        except Exception as e:
            logger.error(f"Error generating lesson plan: {str(e)}")
            raise

@lru_cache()
def get_document_service() -> DocumentService:
    """Get cached document service instance."""
    return DocumentService() 