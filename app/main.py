from fastapi import FastAPI, HTTPException, Depends, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import RedirectResponse, JSONResponse, FileResponse, Response
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.trustedhost import TrustedHostMiddleware
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded
from slowapi.middleware import SlowAPIMiddleware
from starlette.middleware.base import BaseHTTPMiddleware
from starlette.responses import Response
from starlette.config import Config
import logging
from typing import Optional
import tempfile
import os
from pathlib import Path

from app.core.config import get_settings
from app.services.openai_service import get_openai_service
from app.services.msgraph_service import get_msgraph_service
from app.services.twilio_service import get_twilio_service
from app.services.translation_service import get_translation_service
from app.models.api import (
    TextRequest, DocumentRequest, TokenResponse,
    UserInfoResponse, TextResponse, SMSRequest, SMSResponse,
    TranslationRequest, TranslationResponse,
    TranslatedMessageRequest, TranslatedMessageResponse,
    MessageType
)

# Configure logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

# Initialize rate limiter with empty config (no .env required)
limiter = Limiter(
    key_func=get_remote_address,
    app_config=Config(environ={
        "RATELIMIT_STORAGE_URL": "memory://",
        "RATELIMIT_DEFAULT": "100/minute",
        "RATELIMIT_STRATEGY": "fixed-window"
    })
)

# Initialize FastAPI app
app = FastAPI(title=get_settings().APP_NAME)

# Add rate limiting middleware and exception handler
app.state.limiter = limiter
app.add_middleware(SlowAPIMiddleware)
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

# Security Headers Middleware
class SecurityHeadersMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request: Request, call_next):
        response = await call_next(request)
        response.headers["X-Content-Type-Options"] = "nosniff"
        response.headers["X-Frame-Options"] = "DENY"
        response.headers["X-XSS-Protection"] = "1; mode=block"
        response.headers["Strict-Transport-Security"] = "max-age=31536000; includeSubDomains"
        response.headers["Content-Security-Policy"] = "default-src 'self'; img-src 'self' data:; style-src 'self' 'unsafe-inline'"
        return response

# Add security middlewares
app.add_middleware(SecurityHeadersMiddleware)
app.add_middleware(
    TrustedHostMiddleware,
    allowed_hosts=["faraday-ai.onrender.com", "localhost", "127.0.0.1"]
)

# Add CORS middleware with more restrictive settings
app.add_middleware(
    CORSMiddleware,
    allow_origins=["https://faraday-ai.onrender.com", "http://localhost:8000", "http://127.0.0.1:8000"],
    allow_credentials=True,
    allow_methods=["GET", "POST", "OPTIONS", "HEAD"],
    allow_headers=["*"],
)

# Get the absolute path to the static directory
static_dir = Path(__file__).parent / "static"
images_dir = static_dir / "images"
logger.info(f"Serving static files from: {static_dir}")
logger.info(f"Images directory: {images_dir}")

# List all files in the static and images directories
if static_dir.exists():
    logger.info(f"Static directory contents: {[f.name for f in static_dir.iterdir()]}")
    if images_dir.exists():
        logger.info(f"Images directory contents: {[f.name for f in images_dir.iterdir()]}")
    else:
        logger.error(f"Images directory does not exist: {images_dir}")
else:
    logger.error(f"Static directory does not exist: {static_dir}")

# Mount static files
app.mount("/static", StaticFiles(directory=str(static_dir), html=True), name="static")

@app.get("/", response_class=FileResponse)
async def root():
    """Serve the landing page."""
    index_path = static_dir / "index.html"
    logger.info(f"Serving index from: {index_path}")
    if not index_path.exists():
        logger.error(f"Index file not found at: {index_path}")
        raise HTTPException(status_code=404, detail="Index file not found")
    return FileResponse(str(index_path))

@app.get("/favicon.{ext}")
async def favicon(ext: str):
    """Handle favicon requests for different formats."""
    if ext.lower() in ['ico', 'svg']:
        favicon_path = static_dir / "icons" / f"favicon.{ext}"
        if favicon_path.exists():
            return FileResponse(
                str(favicon_path),
                media_type=f"image/{ext}",
                headers={"Cache-Control": "public, max-age=31536000"}
            )
    return Response(status_code=204)

@app.get("/apple-touch-icon{ext:path}")
async def apple_touch_icon(ext: str = ".png"):
    """Handle Apple touch icon requests."""
    # First try SVG
    icon_path = static_dir / "icons" / "favicon.svg"
    if icon_path.exists():
        return FileResponse(
            str(icon_path),
            media_type="image/svg+xml",
            headers={"Cache-Control": "public, max-age=31536000"}
        )
    return Response(status_code=404)

@app.get("/robots.txt")
async def robots():
    """Serve robots.txt file."""
    robots_path = static_dir / "robots.txt"
    if not robots_path.exists():
        return Response(
            content=(
                "User-agent: *\n"
                "Allow: /\n"
                "Allow: /static/images/\n"
                "Allow: /static/icons/\n"
                "Disallow: /debug/\n"
            ),
            media_type="text/plain"
        )
    return FileResponse(str(robots_path))

@app.get("/debug/files")
@limiter.limit("5/minute")
async def debug_files(request: Request):
    """Debug endpoint to list all files in the static directory."""
    try:
        static_files = []
        if static_dir.exists():
            static_files = [str(f.relative_to(static_dir)) for f in static_dir.rglob("*") if f.is_file()]
        return {
            "static_dir": str(static_dir),
            "static_dir_exists": static_dir.exists(),
            "files": static_files,
            "cwd": os.getcwd()
        }
    except Exception as e:
        logger.exception("Error in debug endpoint")
        return {"error": str(e)}

@app.get("/image/{image_name}")
async def get_image(image_name: str):
    """Direct image serving endpoint for debugging."""
    try:
        image_path = images_dir / image_name
        logger.info(f"Attempting to serve image from: {image_path}")
        if not image_path.exists():
            logger.error(f"Image not found: {image_path}")
            available_images = [f.name for f in images_dir.iterdir() if f.is_file()]
            return JSONResponse(
                status_code=404,
                content={
                    "detail": "Image not found",
                    "requested_path": str(image_path),
                    "available_images": available_images
                }
            )
        return FileResponse(str(image_path))
    except Exception as e:
        logger.exception("Error serving image")
        return JSONResponse(
            status_code=500,
            content={"detail": str(e)}
        )

@app.get("/test")
async def test():
    """Health check endpoint."""
    return {"status": "success", "message": "Service is running"}

@app.get("/login")
async def login(msgraph_service = Depends(get_msgraph_service)):
    """Initiate Microsoft Graph authentication."""
    try:
        auth_url = msgraph_service.get_auth_url()
        logger.debug(f"Generated auth URL: {auth_url}")
        return RedirectResponse(url=auth_url)
    except Exception as e:
        logger.error(f"Login error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/callback")
async def callback(
    code: str,
    msgraph_service = Depends(get_msgraph_service)
) -> TokenResponse:
    """Handle Microsoft Graph authentication callback."""
    try:
        result = await msgraph_service.get_token(code)
        if result["status"] == "error":
            raise HTTPException(status_code=400, detail=result["error"])
        return TokenResponse(**result)
    except Exception as e:
        logger.error(f"Callback error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/me")
async def get_user_info(
    request: Request,
    msgraph_service = Depends(get_msgraph_service)
) -> UserInfoResponse:
    """Get user information from Microsoft Graph."""
    token = request.headers.get("Authorization")
    if not token:
        raise HTTPException(status_code=401, detail="No token provided")
    
    result = await msgraph_service.get_user_info(token.split()[1])
    if result["status"] == "error":
        raise HTTPException(status_code=400, detail=result["error"])
    return UserInfoResponse(**result)

@app.post("/generate-text")
async def generate_text(
    request: TextRequest,
    openai_service = Depends(get_openai_service)
) -> TextResponse:
    """Generate text using OpenAI."""
    result = await openai_service.generate_text(
        request.prompt,
        request.structured_output
    )
    if result["status"] == "error":
        raise HTTPException(status_code=500, detail=result["error"])
    return TextResponse(**result)

@app.post("/generate-document")
async def generate_document(request: DocumentRequest) -> FileResponse:
    """Generate a document in the specified format."""
    try:
        # Create a temporary file with the generated content
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{request.output_format}") as temp_file:
            if request.output_format == "docx":
                from docx import Document
                doc = Document()
                doc.add_heading(request.title, 0)
                doc.add_paragraph(request.content)
                doc.save(temp_file.name)
            else:
                raise HTTPException(status_code=400, detail=f"Unsupported format: {request.output_format}")

        return FileResponse(
            temp_file.name,
            media_type="application/octet-stream",
            filename=f"{request.title}.{request.output_format}"
        )
    except Exception as e:
        logger.error(f"Error generating document: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        # Clean up the temporary file
        if 'temp_file' in locals():
            os.unlink(temp_file.name)

@app.post("/send-sms")
async def send_sms(
    request: SMSRequest,
    twilio_service = Depends(get_twilio_service)
) -> SMSResponse:
    """Send an SMS message using Twilio."""
    result = await twilio_service.send_sms(
        request.to_number,
        request.message
    )
    if result["status"] == "error":
        raise HTTPException(status_code=500, detail=result["error"])
    return SMSResponse(**result)

@app.post("/translate")
async def translate_text(
    request: TranslationRequest,
    translation_service = Depends(get_translation_service)
) -> TranslationResponse:
    """Translate text using Google Cloud Translation API."""
    result = await translation_service.translate_text(
        request.text,
        request.target_language,
        request.source_language
    )
    if result["status"] == "error":
        raise HTTPException(status_code=500, detail=result["error"])
    return TranslationResponse(**result)

@app.post("/send-translated-message")
async def send_translated_message(
    request: TranslatedMessageRequest,
    translation_service = Depends(get_translation_service),
    twilio_service = Depends(get_twilio_service)
) -> TranslatedMessageResponse:
    """Translate and send a message via SMS or voice call."""
    try:
        # First, translate the message
        translation_result = await translation_service.translate_text(
            request.message,
            request.target_language,
            request.source_language
        )
        
        if translation_result["status"] == "error":
            raise HTTPException(status_code=500, detail=translation_result["error"])
        
        translated_text = translation_result["translated_text"]
        
        # Then, send the translated message
        if request.message_type == MessageType.SMS:
            delivery_result = await twilio_service.send_sms(
                request.to_number,
                translated_text
            )
            if delivery_result["status"] == "error":
                raise HTTPException(status_code=500, detail=delivery_result["error"])
            
            return TranslatedMessageResponse(
                status="success",
                original_text=request.message,
                translated_text=translated_text,
                message_type="sms",
                delivery_status=delivery_result["status"],
                message_sid=delivery_result.get("message_sid")
            )
        else:  # Voice call
            delivery_result = await twilio_service.make_call(
                request.to_number,
                translated_text,
                f"{request.target_language}-{request.target_language.upper()}"
            )
            if delivery_result["status"] == "error":
                raise HTTPException(status_code=500, detail=delivery_result["error"])
            
            return TranslatedMessageResponse(
                status="success",
                original_text=request.message,
                translated_text=translated_text,
                message_type="voice",
                delivery_status=delivery_result["status"],
                call_sid=delivery_result.get("call_sid")
            )
            
    except Exception as e:
        logger.error(f"Error in send_translated_message: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e)) 